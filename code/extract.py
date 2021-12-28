# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
import os
import sys
import subprocess
t = sys.argv[1]
c = '../results/batch/' + str(t) + '/'
cmd = 'touch %s' % (c + 'results.txt')
subprocess.call(cmd, shell=True)

document = Document()
ml = t.split('/')[1]
if t.split('/')[0] == 'DNA':
    document.add_heading('The results of DNA analysis', 0)
    if ml == 'svm':
        table = document.add_table(rows=7, cols=9)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "C"
        hdr_cells[7].text = "g"
        hdr_cells[8].text = "P"
    elif ml in ['rf']:
        table = document.add_table(rows=7, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "T/K"
        hdr_cells[7].text = "P"
    else:
        table = document.add_table(rows=7, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[7].text = "P"
elif t.split('/')[0] == 'RNA':
    document.add_heading('The results of RNA analysis', 0)
    if ml == 'svm':
        table = document.add_table(rows=6, cols=9)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "C"
        hdr_cells[7].text = "g"
        hdr_cells[8].text = "P"
    elif ml in ['rf']:
        table = document.add_table(rows=6, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "T"
        hdr_cells[7].text = "P"
    else:
        table = document.add_table(rows=6, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[7].text = "P"
elif t.split('/')[0] == 'Protein':
    document.add_heading('The results of protein analysis', 0)
    if ml == 'svm':
        table = document.add_table(rows=13, cols=9)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "C"
        hdr_cells[7].text = "g"
        hdr_cells[8].text = "P"
    elif ml in ['rf']:
        table = document.add_table(rows=13, cols=8)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "T"
        hdr_cells[7].text = "P"
    else:
        table = document.add_table(rows=13, cols=7)
        hdr_cells = table.rows[0].cells
        hdr_cells[1].text = "Acc"
        hdr_cells[2].text = "MCC"
        hdr_cells[3].text = "AUC"
        hdr_cells[4].text = "Sn"
        hdr_cells[5].text = "Sp"
        hdr_cells[6].text = "P"


with open(c + 'results.txt') as f:
    lines = f.readlines()
    for line, i in zip(lines, list(range(len(lines)))):
        line = line.strip().split(' ')
        hdr_cells = table.rows[i + 1].cells
        hdr_cells[0].text = line[0]
        hdr_cells[1].text = str(line[1])
        hdr_cells[2].text = str(line[2])
        hdr_cells[3].text = str(line[3])
        hdr_cells[4].text = str(line[4])
        hdr_cells[5].text = str(line[5])
        if t.split('/')[1] == 'svm':
            hdr_cells[6].text = str(line[6])
            hdr_cells[7].text = str(line[7])
            hdr_cells[8].text = str(line[8])
        elif t.split('/')[1] in ['rf']:
            hdr_cells[6].text = str(line[6])
            hdr_cells[7].text = str(line[7])
        else:
            hdr_cells[6].text = str(line[6])

document.save(c + 'results.docx')
os.remove(c + 'results.txt')
